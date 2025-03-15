import skills
import yaml
from math import ceil, floor
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Mm, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from paragraph import paragraph

class CVBldr:
    def __init__(self, expFile):

        with open('header.yaml', 'r') as file:
            self.headers = yaml.safe_load(file)
            file.close()

        with open('formatting.yaml', 'r') as file:
            self.formatting = yaml.safe_load(file)
            file.close()

        self.createDoc()
        self.setMargins()
        self.getExpFile(expFile)

    def createDoc(self):
        self.doc = Document()

    def getExpFile(self, expFile):
        
        with open(expFile) as f:
            txt = f.read()
            f.close()
        
        # add the formatting to the string. eval() does the reverse
        txt = repr(txt)

        # Remove comments denoted by '%%'. This is done by splitting on %% and recovering only odd splits. 
        # Obviously this will break if any comments are not closed, but this would be extremely obvious in Obsidian
        txt = ''.join(txt.split('%%')[::2])

        self.expList = []
        exps =  txt.split('### ')
        # Get rid of the header
        exps.pop(0)

        for exp in exps:
            # Split by newline
            expDetails = exp.split(r'\n')
            
            # Don't handle the skills section
            if 'Skills' in expDetails[0]:
                self.skills = [x.strip('- ') for x in expDetails if '-' in x[0:3]]
            else:

                # Remove if empty (happens after multi-newlines)
                expDetails = [x for x in expDetails if x]

                # remove leading and trailing spaces
                expDetails = [x.strip(' ') for x in expDetails if x]

                expDict = {
                    'company' : expDetails[0],
                    # get the element with startDate, split off the 'startdate:' and then strip spaces
                    'startDate' : [x for x in expDetails if 'startDate' in x][0].split(':')[-1].strip(' '),
                    'endDate' : [x for x in expDetails if 'endDate' in x][0].split(':')[-1].strip(' '),
                    'title' : [x for x in expDetails if 'title' in x][0].split(':')[-1].strip(' '),
                    'accolades' : [x.strip('- ') for x in expDetails if '-' in x[0:3] and x.strip('- ')],
                }
            
                self.expList.append(expDict)
    
    def addHeader(self):
        
        if self.headers['name']:
            self.getParagraph(self.headers['name'], ['defaultTitleFormat', 'titleFormat'])
        
        if self.headers['roles']:
            roles =  [' | '.join(self.headers['roles'])]
            self.getParagraph(roles, ['defaultTitleFormat', 'rolesFormat'])

        if self.headers['contact'].values():
            contactList = ['   |   '.join(self.headers['contact'].values())]
            self.getParagraph(contactList, ['defaultTitleFormat', 'contactFormat'])

    def addSkills(self):

        # Skills Header
        self.getParagraph('Skills', ['defaultTitleFormat', 'headerFormat'])

        rows = ceil(len(self.skills)/2)
        table = self.doc.add_table(rows=0, cols=2)
        table.style.paragraph_format.space_after = Pt(0)
        for row in range(0, rows):
            row_cells = table.add_row().cells
            if self.skills:
                cell1 = self.getParagraph(self.skills.pop(0), ['bulletFormat'], row_cells[0].paragraphs[0])
            else:   
                cell1 = self.getParagraph('', [], row_cells[0].paragraphs[0])
            if self.skills:
                cell2 = self.getParagraph(self.skills.pop(0), ['bulletFormat'], row_cells[1].paragraphs[0])
            else:   
                cell2 = self.getParagraph('', [], row_cells[1].paragraphs[0])

            if row == (rows-1):
                cell1.para.paragraph_format.space_after = Pt(0)
                cell2.para.paragraph_format.space_after = Pt(0)


    def setMargins(self, section=0, marginName='docMargins'):
        self.doc.sections[section].top_margin = Mm(self.formatting[marginName][0])
        self.doc.sections[section].right_margin = Mm(self.formatting[marginName][1])
        self.doc.sections[section].bottom_margin = Mm(self.formatting[marginName][2])
        self.doc.sections[section].left_margin = Mm(self.formatting[marginName][3])

    def saveDoc(self):
        self.doc.save('New Rezz.docx')

    def getParagraph(self, runs, formats, existing_paragraph=None):
    
        info = self.formatting['defaultFormat'].copy()
        for format in formats:
            info.update(self.formatting[format])

        if isinstance(runs, str):
            runs = [runs]

        info.update({'runs' : runs})
        return paragraph(self.doc, info, existing_paragraph)

    def addSpacer(self):
        spacer = self.formatting['defaultFormat'].copy()
        spacer.update(self.formatting['spacerFormat'])
        spacer.update({'runs' : ''})
        spacer = paragraph(self.doc, spacer)

    def addEducation(self):
                
        # Edu Header
        self.getParagraph('Education', ['defaultTitleFormat', 'headerFormat'])

        education = self.headers['education']
        degrees = []
        for degree in education:
            degreeString = f"{education[degree]['diploma']} from the {education[degree]['school']}, completed {education[degree]['completionDate']}"
            degrees.append(degreeString)
        self.getParagraph(degrees, ['bulletFormat'])

    def addExperience(self):

        # Professional Experience Header
        self.getParagraph('Professional Experience', ['defaultTitleFormat', 'headerFormat'])

        for exp in self.expList:
            table = self.doc.add_table(rows=0, cols=2)
            table.style.paragraph_format.space_after = Pt(0)
            row_cells = table.add_row().cells
            title = self.getParagraph(exp['title'], ['defaultTitleFormat'], row_cells[0].paragraphs[0])
            tenureStr = self.getTenureStr(exp)
            tenurePara = self.getParagraph(tenureStr, ['defaultTitleFormat'], row_cells[1].paragraphs[0])
            tenurePara.para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            self.getParagraph(exp['company'], [])
            for accolade in exp['accolades']:
                self.getParagraph(accolade, ['bulletFormat'])
            self.addSpacer()

    def getTenureStr(self, expDict):

        startDate = datetime.strptime(expDict['startDate'], '%B %Y') if expDict['startDate'] else datetime.today()
        endDate = datetime.strptime(expDict['endDate'], '%B %Y') if expDict['endDate'] else datetime.today()
        year = floor((endDate-startDate).days / 365)
        monthDelta = (endDate.month-startDate.month)
        month = monthDelta if monthDelta >= 0 else monthDelta + 12

        if expDict['endDate']:
            endDateStr = endDate.strftime('%B %Y')
        else:
            endDateStr = 'Present'

        # I don't see a way around manually coding how to handle different lengths of time
        if year == 0:
            yearStr = f''
        elif year == 1:
            yearStr = f'{year} year, '
        else:
            yearStr = f'{year} years, '

        if month == 0:
            monthStr = f''
        elif month == 1:
            monthStr = f'{month} month'
        else:
            monthStr = f'{month} months' 
        
        tenureStr = f"{startDate.strftime('%B %Y')} - {endDateStr} ({yearStr}{monthStr})"
        return tenureStr

if __name__ == '__main__':
    config = yaml.safe_load(open('config.yaml'))
    expFile = eval(config['expFile']).replace('\\', '/')
    resume = CVBldr(expFile)
    resume.addHeader()
    resume.addSpacer()
    resume.addSkills()
    resume.addSpacer()
    resume.addEducation()
    resume.addSpacer()
    resume.addExperience()
    resume.saveDoc()
 